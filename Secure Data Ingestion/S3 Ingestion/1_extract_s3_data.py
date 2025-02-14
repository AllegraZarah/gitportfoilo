import os
import boto3
import pandas as pd
import psycopg2
from docx import Document
from io import BytesIO
from pypdf import PdfReader
from dotenv import load_dotenv

def connect_to_database():
    """ 
    Establishes a connection to a PostgreSQL database.

    This function attempts to establish a connection to a PostgreSQL database using the psycopg2 library.
    The database name, username, password, host, and port are hardcoded in the function. 

    Required:
    - import psycopg2
    - import os
    - from dotenv import load_dotenv
    - .env file in project directory

    Parameters:
    All parameters are harvested from a .env file attached to the project.
    Note that the variables should be defined as stated below in brackets.
    - host (str) ["POSTGRES_ADDRESS"]: The host of the database.
    - dbname (str) ["DATABASE"]: The name of the database.
    - user (str) ["POSTGRES_USERNAME"]: The username used to authenticate with the database.
    - password (str) ["POSTGRES_PASSWORD"]: The password used to authenticate with the database.
    - port (str) ["POSTGRES_PORT"]: The port of the database.

    Returns:
    connection (psycopg2.extensions.connection): A connection object that represents the database connection. 
    If the connection is successful, this object can be used to perform SQL operations on the database. 
    If the connection fails, the function will print an error message and return None.

    Raises:
    Exception: An exception is raised if there is an error when trying to connect to the database. 
    The exception message will contain information about the error.

    """
     
    try:
        # Load the environment variables from the .env file
        load_dotenv()

        # Get the connection details from the environment variables
        pg_datasource_host = os.getenv("POSTGRES_ADDRESS")
        pg_datasource_dbname = os.getenv("DATABASE")
        pg_datasource_username = os.getenv("POSTGRES_USERNAME")
        pg_datasource_pw = os.getenv("POSTGRES_PASSWORD")
        pg_datasource_port = os.getenv("POSTGRES_PORT")

        # Establish a connection to the database
        connection = psycopg2.connect(database = pg_datasource_dbname,
                                       user = pg_datasource_username,
                                       host = pg_datasource_host,
                                       password = pg_datasource_pw,
                                       port = 5432)
        print("Connection to Database Successful")
        
        return connection

    except Exception as e:
        print("An error occurred:", e)
        return None

def connect_to_s3_bucket(bucket_name):
    """ 
    Establishes a connection to a PostgreSQL database.

    This function attempts to establish a connection to a PostgreSQL database using the psycopg2 library.
    The database name, username, password, host, and port are hardcoded in the function. 

    Required:
    - import os
    - import boto3

    Parameters:
    The first three(3) parameters are harvested from a .env file attached to the project.
    Note that the variables should be defined as stated below in brackets.
    The last function should be entered at the point of calling the function.
    - aws_access_key_id (str) ["ACCESS_KEY"]: The AWS access key ID for your account.
    - aws_secret_access_key (str) ["SECRET_KEY"]: The AWS secret access key for your account.
    - region_name (str) ["REGION"]: The name of the AWS region where the S3 bucket is located.
    - bucket_name (str): The name of the S3 bucket you want to connect to.

    Returns:
    s3 (boto3.resources.factory.s3.ServiceResource): A resource representing Amazon S3. 
    If the connection is successful, this object can be used to interact with the contents of the s3 bucket. 
    If the connection fails, the function will print an error message and return None.

    Raises:
    Exception: An exception is raised if there is an error when trying to connect to the s3 bucket. 
    The exception message will contain information about the error.

    """

    try:
        # Get the connection details from the environment variables
        aws_access_key_id = os.getenv('ACCESS_KEY')
        aws_secret_access_key = os.getenv('SECRET_KEY')
        region_name = os.getenv('REGION')

        # Create a boto3 session to the s3 bucket of concern
        s3 = boto3.resource('s3',
                            aws_access_key_id=aws_access_key_id,
                            aws_secret_access_key=aws_secret_access_key,
                            region_name=region_name
                            )
        
        # Test the connection by trying to access the bucket
        s3.meta.client.head_bucket(Bucket=bucket_name)
        
        print("Connection to S3 successful")

        return s3
        

    except Exception as e:
        print("An error occurred:", e)

def ensure_table_exists(connection, schema_name, table_name):
    """ 
    Ensures that a specified table exists within a given schema in a PostgreSQL database.

    This function checks whether a table exists in the specified schema of the PostgreSQL database. 
    If the table does not exist, the function creates the schema (if it doesn't already exist) and then 
    creates the specified table with predefined columns.

    Required:
    - Function: connect_to_database
    - import psycopg2

    Parameters:
    - connection (psycopg2.extensions.connection): The database connection object used to interact with the PostgreSQL database.
    - schema_name (str): The name of the schema where the table should be checked or created.
    - table_name (str): The name of the table to check or create within the specified schema.

    Returns:
    None.
    The function prints a message indicating whether the table was successfully created or if it already exists.

    Raises:
    Exception: An exception might be raised if there are issues with the SQL execution or database connection, 
    though the function itself does not explicitly handle exceptions beyond printing basic status messages.

    Notes:
    - The table created by this function has the following columns: 
      - file_hash (varchar(64))
      - document_name (varchar(64))
      - creation_date (TIMESTAMP)
      - document_text (TEXT)
    - The function uses SQL statements executed through a cursor object, and it commits transactions to ensure changes persist.

    """
    schema_name = schema_name
    table_name = table_name

    # Open a cursor to perform database operations
    cursor = connection.cursor()

    # Execute a command
    cursor.execute("SELECT EXISTS(SELECT 1 FROM information_schema.tables WHERE table_schema = %s AND table_name = %s);", (schema_name, table_name))

    # fetch the result
    table_exists = cursor.fetchone()[0]


    if not table_exists:
        with connection.cursor() as cursor:
            # Create the schema if it doesn't exist
            cursor.execute(f"CREATE SCHEMA IF NOT EXISTS {schema_name};")
            # Commit execution to ensure the changes to the database persist
            connection.commit()

            # Create the table
            # Note we are using {} inplace of %s here because psycopg2 treats %s as a string literal and will put quotes around the texts.
            # That works for filtering, but will not work for the fullpath name of the table.
            cursor.execute(f""" CREATE TABLE {schema_name}.{table_name} (file_hash varchar(64), document_name varchar(64), creation_date TIMESTAMP, document_text TEXT)""")
            # Commit second execution to ensure the changes to the database persist
            connection.commit()
        print(f"Table:{table_name} successfully created in Schema: {schema_name}")

    elif table_exists:
        print("Table already exists within schema")

    else:
        print("Something is not right and should be checked out")


# %%
def extract_text_from_s3_files(file, cv_file):
    """
    Extracts text content from PDF or DOCX files stored in an S3 bucket.

    This function handles the extraction of text from files retrieved from an S3 bucket. Depending on the file type (PDF or DOCX), 
    it uses the appropriate libraries to extract the text content.

    Required:
    - import docx: For extracting text from DOCX files using python-docx.
    - import io.BytesIO: For handling in-memory file operations.
    - import pypdf.PdfReader: For reading and extracting text from PDF files.

    Parameters:
    - file (boto3.s3.Object): The S3 file object that represents the file stored in the bucket. 
      The function checks the file extension to determine whether it's a PDF or DOCX.
    - cv_file (bytes): The content of the file retrieved from S3, passed as a byte stream for processing.

    Returns:
    - pagewise_content (str): A string containing the extracted text content from the file. If an error occurs during extraction, 
      the function returns a string indicating that no text could be extracted, along with the filename.

    Raises:
    Exception: An exception is handled within the function to catch errors during file processing, such as issues with reading the file format.

    Notes:
    - For PDF files: The function uses PyPDF to read and extract text from each page.
    - For DOCX files: The function uses python-docx to read and extract text from each paragraph.
    - The function currently does not calculate the number of pages for DOCX files, as python-docx does not provide this capability.
    - If the file type is unsupported or an error occurs, a message indicating the issue is returned.

    """
    # Create a temporary string to store the text from all pages of a file
    pagewise_content = ''

    try:
            
        if file.key.endswith('.pdf'):
            # Read the PDF file using PyPDF
            pdf_reader = PdfReader(BytesIO(cv_file))

            # Get the number of pages of the file
            no_of_pages = len(pdf_reader.pages)

            # Extract text from each page and append it to the list
            for page in range(no_of_pages):
                cv_pages = pdf_reader.pages[page]
                pagewise_content += cv_pages.extract_text()

        elif file.key.endswith('.docx'):
            # Read the .docx file using python-docx
            doc_reader = Document(BytesIO(cv_file))

            # Get the number of pages of the file
            # The python-docx library lacks a built-in method to get the number of pages in a Word document. 
            # This is because page count can vary based on page size, font size, and formatting settings, 
            # which are not stored in the .docx file.
            # Hence, page count for .docx will be none. This is a reason why .pdf files should be the only allowed file type
            no_of_pages = None

            # Extract text from each paragraph and print it
            for paragraph in doc_reader.paragraphs:
                pagewise_content += paragraph.text


    except Exception as e:
        pagewise_content = f"No Text In CV With Filename: {file.key}"

    return pagewise_content

# %%
def process_documents(s3_connection_function, bucket_name, database_connection_function, schema_name, table_name):
    """
    Adds unique text content extracted from files in an S3 bucket to a PostgreSQL database.

    This is the main function in this script
    This function connects to an S3 bucket, iterates through all files in the bucket, and checks if the content of each file (based on a unique hash) 
    already exists in the specified database table. If the content is unique, it extracts the text from the file and inserts it into the database.

    Required:
    - Function: s3_connection_function - Function to establish a connection to the S3 bucket.
    - Function: database_connection_function - Function to establish a connection to the PostgreSQL database.
    - Function: extract_text_from_s3_files - Function to extract text content from PDF or DOCX files.

    Parameters:
    - s3_connection_function (function): A function that returns a connected S3 resource.
    - bucket_name (str): The name of the S3 bucket containing the files.
    - database_connection_function (psycopg2.extensions.connection): A function that returns a connection to the PostgreSQL database.
    - schema_name (str): The name of the schema where the table resides.
    - table_name (str): The name of the table to check for existing records and insert new ones.

    Returns:
    - None: The function does not return a value, but it prints messages indicating the status of the operations performed.

    Process:
    1. The function connects to the specified S3 bucket and iterates through each file.
    2. For each file, it calculates a unique hash (ETag) and checks if this hash already exists in the database.
    3. If the hash is not found in the database, the function extracts text from the file and inserts the content along with metadata into the specified database table.
    4. The function also queries and displays the last 5 records from the table, ordered by creation date.

    Notes:
    - The function assumes that the database table has columns named 'file_hash', 'document_name', 'creation_date', and 'document_text'.
    - If a file's content already exists in the database (based on the hash), it is skipped.
    - The function commits the changes to the database after each insertion. 

    """
    bucket = s3_connection_function.Bucket(bucket_name)
    connection = database_connection_function

    ensure_table_exists(connection, schema_name, table_name)

    for file in bucket.objects.all():
        cv_file_hash = file.e_tag.replace('"', "")

        with connection.cursor() as cursor:
            cursor.execute("""SELECT EXISTS(SELECT 1 FROM {}.{} WHERE {} = '{}');""".format(schema_name, table_name, "file_hash", cv_file_hash))

            record_exists = cursor.fetchone()[0]
            # print(record_exists)

            if not record_exists:
                # Extract text from document
                # Get the PDF file from the S3 bucket
                cv_file = s3_connection_function.Object(bucket_name, file.key).get()['Body'].read()
                cv_title = file.key
                updated_at = file.last_modified
                pagewise_content = extract_text_from_s3_files(file, cv_file)

                 # Insert data into the database
                insert_query = """INSERT INTO {}.{} (file_hash, document_name, creation_date, document_text) 
                                    VALUES (%s, %s, %s, %s) RETURNING *;""".format(schema_name, table_name)
                cursor.execute(insert_query, (cv_file_hash, cv_title, updated_at, pagewise_content))
                
                
                # Make the changes to the database persistent
                connection.commit()
                print(f"File content successfully added to the table: {table_name} in the schema: {schema_name}")
            
            elif record_exists:
                continue

    # View the table to confirm output
    query = f"SELECT * FROM {schema_name}.{table_name} ORDER BY creation_date DESC LIMIT 5"
    return pd.read_sql_query(query, connection)

if __name__ == "__main__":
    s3 = connect_to_s3_bucket('bucket-name')
    db = connect_to_database()
    if s3 and db:
        process_documents(s3, 'bucket-name', db, 'documents', 'extracted_content')